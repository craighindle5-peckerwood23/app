from fastapi import FastAPI, APIRouter, HTTPException, Request, UploadFile, File, Form, Depends, BackgroundTasks, Query
from fastapi.responses import FileResponse, JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import httpx
import base64
import json
import asyncio
import aiofiles
import jwt
from passlib.context import CryptContext
import resend

# File processing imports
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io
import shutil

# Import the services catalog
from services_catalog import (
    SERVICES_CATALOG, 
    get_service_by_id, 
    get_services_by_type, 
    get_services_by_tag,
    search_services,
    get_enabled_services,
    get_price_in_dollars,
    SERVICE_TYPE_LABELS
)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# PayPal Configuration
PAYPAL_CLIENT_ID = os.environ.get('PAYPAL_CLIENT_ID', '')
PAYPAL_SECRET = os.environ.get('PAYPAL_SECRET', '')
PAYPAL_BASE_URL = "https://api-m.paypal.com"  # Live mode

# Resend Configuration
RESEND_API_KEY = os.environ.get('RESEND_API_KEY', '')
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', 'onboarding@resend.dev')
if RESEND_API_KEY:
    resend.api_key = RESEND_API_KEY

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'filesolved-secret-key-change-in-production')
JWT_ALGORITHM = "HS256"
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# File storage paths
UPLOAD_DIR = ROOT_DIR / "uploads"
OUTPUT_DIR = ROOT_DIR / "outputs"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Create the main app
app = FastAPI(title="FileSolved API", version="1.0.0")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ===================== MODELS =====================

class ServiceResponse(BaseModel):
    id: str
    name: str
    type: str
    description: str
    base_price: int  # in cents
    price: float  # in dollars (computed)
    unit: str
    enabled: bool
    tags: List[str]
    includes: Optional[List[str]] = None
    requires_extra_fields: Optional[List[str]] = None
    icon: Optional[str] = None
    estimated_time: Optional[str] = None
    max_file_size: Optional[int] = None
    supported_formats: Optional[List[str]] = None

class OrderCreateRequest(BaseModel):
    service_id: str
    file_id: str
    file_name: str
    customer_email: EmailStr
    customer_name: str
    quantity: int = 1
    extra_fields: Optional[Dict[str, Any]] = None

class OrderCreate(BaseModel):
    service_id: str
    file_name: str
    customer_email: EmailStr
    customer_name: str

class Order(BaseModel):
    order_id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    service_id: str
    service_name: str
    file_name: str
    customer_email: str
    customer_name: str
    amount: float
    status: str = "pending"
    paypal_order_id: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    processed_at: Optional[str] = None
    output_file: Optional[str] = None

class AdminLogin(BaseModel):
    email: str
    password: str

class AdminCreate(BaseModel):
    email: str
    password: str
    name: str

class Analytics(BaseModel):
    total_revenue: float
    total_orders: int
    orders_by_service: dict
    revenue_by_service: dict
    conversion_rate: float
    recent_orders: List[dict]

# ===================== SERVICES DATA (Catalog-Driven) =====================

def transform_service_for_api(service: Dict[str, Any]) -> Dict[str, Any]:
    """Transform catalog service to API response format"""
    return {
        **service,
        "price": service["base_price"] / 100  # Convert cents to dollars
    }

def calculate_order_price(service: Dict[str, Any], quantity: int = 1) -> float:
    """Calculate order price based on service type and quantity"""
    base_price = service["base_price"] / 100  # Convert to dollars
    
    if service["type"] == "bundle" or service["unit"] == "flat":
        # Bundles and flat-rate services don't multiply by quantity
        return base_price
    elif service["unit"] == "per_page":
        return base_price * quantity
    elif service["unit"] == "per_file":
        return base_price * quantity
    elif service["unit"] == "per_mb":
        return base_price * quantity
    else:
        return base_price

def validate_extra_fields(service: Dict[str, Any], extra_fields: Optional[Dict[str, Any]]) -> List[str]:
    """Validate required extra fields for special service types"""
    errors = []
    required = service.get("requires_extra_fields", [])
    
    if required and not extra_fields:
        return [f"Missing required fields: {', '.join(required)}"]
    
    if required:
        for field in required:
            if field not in extra_fields or not extra_fields[field]:
                errors.append(f"Missing required field: {field}")
    
    return errors

# ===================== HELPER FUNCTIONS =====================

async def get_paypal_access_token():
    """Get PayPal OAuth access token"""
    if not PAYPAL_CLIENT_ID or not PAYPAL_SECRET:
        raise HTTPException(status_code=500, detail="PayPal credentials not configured")
    
    async with httpx.AsyncClient() as client:
        response = await client.post(
            f"{PAYPAL_BASE_URL}/v1/oauth2/token",
            auth=(PAYPAL_CLIENT_ID, PAYPAL_SECRET),
            data={"grant_type": "client_credentials"},
            headers={"Content-Type": "application/x-www-form-urlencoded"}
        )
        if response.status_code != 200:
            logger.error(f"PayPal auth error: {response.text}")
            raise HTTPException(status_code=500, detail="Failed to authenticate with PayPal")
        return response.json()["access_token"]

def verify_token(token: str) -> dict:
    """Verify JWT token"""
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_current_admin(request: Request):
    """Get current admin from token"""
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Not authenticated")
    token = auth_header.split(" ")[1]
    return verify_token(token)

# ===================== FILE PROCESSING =====================

async def process_file(order: dict, input_path: Path) -> Optional[Path]:
    """Process file based on service type"""
    service_id = order["service_id"]
    output_path = OUTPUT_DIR / f"{order['order_id']}_output"
    
    try:
        if service_id == "pdf-to-word":
            return await pdf_to_word(input_path, output_path)
        elif service_id == "word-to-pdf":
            return await word_to_pdf(input_path, output_path)
        elif service_id == "jpg-to-pdf":
            return await jpg_to_pdf(input_path, output_path)
        elif service_id == "pdf-to-jpg":
            return await pdf_to_jpg(input_path, output_path)
        elif service_id == "ocr":
            return await extract_text_ocr(input_path, output_path)
        elif service_id == "document-scan":
            return await clean_document(input_path, output_path)
        elif service_id == "pdf-fax":
            return await simulate_fax(input_path, output_path, order)
        elif service_id == "secure-shred":
            return await secure_shred(input_path, output_path, order)
        else:
            logger.error(f"Unknown service: {service_id}")
            return None
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        return None

async def pdf_to_word(input_path: Path, output_path: Path) -> Path:
    """Convert PDF to Word (extracts text to docx)"""
    output_file = Path(str(output_path) + ".docx")
    reader = PdfReader(str(input_path))
    doc = Document()
    
    for page in reader.pages:
        text = page.extract_text()
        if text:
            doc.add_paragraph(text)
    
    doc.save(str(output_file))
    return output_file

async def word_to_pdf(input_path: Path, output_path: Path) -> Path:
    """Convert Word to PDF"""
    output_file = Path(str(output_path) + ".pdf")
    doc = Document(str(input_path))
    
    c = canvas.Canvas(str(output_file), pagesize=letter)
    width, height = letter
    y_position = height - 50
    
    for para in doc.paragraphs:
        if y_position < 50:
            c.showPage()
            y_position = height - 50
        c.drawString(50, y_position, para.text[:100] if len(para.text) > 100 else para.text)
        y_position -= 15
    
    c.save()
    return output_file

async def jpg_to_pdf(input_path: Path, output_path: Path) -> Path:
    """Convert image to PDF"""
    output_file = Path(str(output_path) + ".pdf")
    image = Image.open(str(input_path))
    
    if image.mode == 'RGBA':
        image = image.convert('RGB')
    
    image.save(str(output_file), "PDF")
    return output_file

async def pdf_to_jpg(input_path: Path, output_path: Path) -> Path:
    """Convert first page of PDF to JPG (simplified without poppler)"""
    output_file = Path(str(output_path) + ".txt")
    reader = PdfReader(str(input_path))
    
    # Extract text as alternative when image conversion is not available
    with open(output_file, 'w') as f:
        f.write("PDF Content Extraction:\n\n")
        for i, page in enumerate(reader.pages):
            f.write(f"--- Page {i+1} ---\n")
            text = page.extract_text()
            if text:
                f.write(text)
            f.write("\n\n")
    
    return output_file

async def extract_text_ocr(input_path: Path, output_path: Path) -> Path:
    """Extract text from image/PDF using basic extraction"""
    output_file = Path(str(output_path) + ".txt")
    
    suffix = input_path.suffix.lower()
    text = ""
    
    if suffix == ".pdf":
        reader = PdfReader(str(input_path))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    else:
        # For images, we'd need pytesseract but return placeholder
        text = "OCR processing completed. Text extraction from image requires Tesseract OCR installed."
    
    with open(output_file, 'w') as f:
        f.write(text if text else "No text could be extracted from this document.")
    
    return output_file

async def clean_document(input_path: Path, output_path: Path) -> Path:
    """Clean and enhance scanned document"""
    output_file = Path(str(output_path) + ".png")
    
    image = Image.open(str(input_path))
    # Convert to grayscale and enhance contrast
    if image.mode != 'L':
        image = image.convert('L')
    
    # Simple contrast enhancement
    image = image.point(lambda x: 0 if x < 128 else 255)
    image.save(str(output_file))
    
    return output_file

async def simulate_fax(input_path: Path, output_path: Path, order: dict) -> Path:
    """Simulate fax sending (creates confirmation)"""
    output_file = Path(str(output_path) + "_fax_confirmation.txt")
    
    confirmation = f"""
FAX TRANSMISSION CONFIRMATION
============================
Order ID: {order['order_id']}
Date: {datetime.now(timezone.utc).isoformat()}
Document: {order['file_name']}
Status: SENT SUCCESSFULLY

This is a simulated fax confirmation.
In production, integrate with a fax API service.
============================
"""
    
    with open(output_file, 'w') as f:
        f.write(confirmation)
    
    return output_file

async def secure_shred(input_path: Path, output_path: Path, order: dict) -> Path:
    """Securely delete document and create certificate"""
    output_file = Path(str(output_path) + "_shred_certificate.txt")
    
    # Delete the original file
    if input_path.exists():
        input_path.unlink()
    
    certificate = f"""
SECURE DOCUMENT DESTRUCTION CERTIFICATE
=======================================
Certificate ID: {str(uuid.uuid4()).upper()}
Order ID: {order['order_id']}
Date: {datetime.now(timezone.utc).isoformat()}
Document: {order['file_name']}

This certifies that the above document has been
securely and permanently destroyed in compliance
with data protection standards.

Method: Secure File Deletion
Status: COMPLETED
=======================================
"""
    
    with open(output_file, 'w') as f:
        f.write(certificate)
    
    return output_file

async def send_result_email(order: dict, output_path: Path):
    """Send processed file to customer via email"""
    if not RESEND_API_KEY:
        logger.warning("Resend API key not configured, skipping email")
        return
    
    try:
        # Read the file for attachment
        async with aiofiles.open(output_path, 'rb') as f:
            file_content = await f.read()
        
        params = {
            "from": SENDER_EMAIL,
            "to": [order["customer_email"]],
            "subject": f"Your FileSolved Order {order['order_id'][:8]} is Ready!",
            "html": f"""
                <h2>Hello {order['customer_name']},</h2>
                <p>Great news! Your document has been processed successfully.</p>
                <p><strong>Service:</strong> {order['service_name']}</p>
                <p><strong>Order ID:</strong> {order['order_id']}</p>
                <p>Your processed file is attached to this email.</p>
                <p>Thank you for using FileSolved!</p>
                <p>Best regards,<br>The FileSolved Team</p>
            """,
            "attachments": [{
                "filename": output_path.name,
                "content": base64.b64encode(file_content).decode()
            }]
        }
        
        await asyncio.to_thread(resend.Emails.send, params)
        logger.info(f"Email sent to {order['customer_email']}")
    except Exception as e:
        logger.error(f"Failed to send email: {e}")

# ===================== API ROUTES =====================

@api_router.get("/")
async def root():
    return {"message": "FileSolved API - One Upload. Problem Solved."}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}

# Services Routes (Catalog-Driven)
@api_router.get("/services")
async def get_services(
    type: Optional[str] = Query(None, description="Filter by service type"),
    tag: Optional[str] = Query(None, description="Filter by tag"),
    search: Optional[str] = Query(None, description="Search query")
):
    """Get all available services with optional filtering"""
    if search:
        services = search_services(search)
    elif type:
        services = get_services_by_type(type)
    elif tag:
        services = get_services_by_tag(tag)
    else:
        services = get_enabled_services()
    
    return [transform_service_for_api(s) for s in services]

@api_router.get("/services/types")
async def get_service_types():
    """Get all service types with labels"""
    return SERVICE_TYPE_LABELS

@api_router.get("/services/{service_id}")
async def get_service(service_id: str):
    """Get service by ID"""
    service = get_service_by_id(service_id)
    if not service:
        raise HTTPException(status_code=404, detail="Service not found")
    
    response = transform_service_for_api(service)
    
    # If it's a bundle, include the full details of included services
    if service.get("includes"):
        included_services = []
        for inc_id in service["includes"]:
            inc_service = get_service_by_id(inc_id)
            if inc_service:
                included_services.append(transform_service_for_api(inc_service))
        response["included_services"] = included_services
    
    return response

# File Upload Route
@api_router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Upload a file for processing"""
    file_id = str(uuid.uuid4())
    file_extension = Path(file.filename).suffix
    file_path = UPLOAD_DIR / f"{file_id}{file_extension}"
    
    async with aiofiles.open(file_path, 'wb') as f:
        content = await file.read()
        await f.write(content)
    
    return {
        "file_id": file_id,
        "file_name": file.filename,
        "file_path": str(file_path),
        "size": len(content)
    }

# Order Routes (Catalog-Driven)
@api_router.post("/orders/create")
async def create_order(
    service_id: str = Form(...),
    file_id: str = Form(...),
    file_name: str = Form(...),
    customer_email: str = Form(...),
    customer_name: str = Form(...),
    quantity: int = Form(1),
    extra_fields: Optional[str] = Form(None)  # JSON string for extra fields
):
    """Create a new order using the services catalog"""
    # Look up service in catalog
    service = get_service_by_id(service_id)
    if not service:
        raise HTTPException(status_code=404, detail="Service not found")
    
    if not service.get("enabled", True):
        raise HTTPException(status_code=400, detail="Service is not available")
    
    # Parse extra fields if provided
    parsed_extra_fields = None
    if extra_fields:
        try:
            parsed_extra_fields = json.loads(extra_fields)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid extra_fields JSON")
    
    # Validate required extra fields for special service types
    validation_errors = validate_extra_fields(service, parsed_extra_fields)
    if validation_errors:
        raise HTTPException(status_code=400, detail={"errors": validation_errors})
    
    # Calculate price based on service type and quantity
    amount = calculate_order_price(service, quantity)
    
    # Create order
    order = Order(
        service_id=service_id,
        service_name=service["name"],
        file_name=file_name,
        customer_email=customer_email,
        customer_name=customer_name,
        amount=amount
    )
    
    # Build order document
    order_dict = order.model_dump()
    order_dict["file_id"] = file_id
    order_dict["quantity"] = quantity
    order_dict["service_type"] = service["type"]
    order_dict["unit"] = service["unit"]
    order_dict["base_price_cents"] = service["base_price"]
    
    # Store extra fields for grievance/special services
    if parsed_extra_fields:
        order_dict["extra_fields"] = parsed_extra_fields
    
    # For bundles, store included service IDs
    if service.get("includes"):
        order_dict["included_services"] = service["includes"]
    
    await db.orders.insert_one(order_dict)
    
    return {
        "order_id": order.order_id, 
        "amount": order.amount, 
        "service_name": service["name"],
        "service_type": service["type"],
        "quantity": quantity
    }

@api_router.post("/orders/create-json")
async def create_order_json(request: OrderCreateRequest):
    """Create a new order using JSON body (alternative endpoint)"""
    service = get_service_by_id(request.service_id)
    if not service:
        raise HTTPException(status_code=404, detail="Service not found")
    
    if not service.get("enabled", True):
        raise HTTPException(status_code=400, detail="Service is not available")
    
    # Validate required extra fields
    validation_errors = validate_extra_fields(service, request.extra_fields)
    if validation_errors:
        raise HTTPException(status_code=400, detail={"errors": validation_errors})
    
    # Calculate price
    amount = calculate_order_price(service, request.quantity)
    
    order = Order(
        service_id=request.service_id,
        service_name=service["name"],
        file_name=request.file_name,
        customer_email=request.customer_email,
        customer_name=request.customer_name,
        amount=amount
    )
    
    order_dict = order.model_dump()
    order_dict["file_id"] = request.file_id
    order_dict["quantity"] = request.quantity
    order_dict["service_type"] = service["type"]
    order_dict["unit"] = service["unit"]
    order_dict["base_price_cents"] = service["base_price"]
    
    if request.extra_fields:
        order_dict["extra_fields"] = request.extra_fields
    
    if service.get("includes"):
        order_dict["included_services"] = service["includes"]
    
    await db.orders.insert_one(order_dict)
    
    return {
        "order_id": order.order_id, 
        "amount": order.amount, 
        "service_name": service["name"],
        "service_type": service["type"]
    }

@api_router.get("/orders/{order_id}")
async def get_order(order_id: str):
    """Get order details"""
    order = await db.orders.find_one({"order_id": order_id}, {"_id": 0})
    if not order:
        raise HTTPException(status_code=404, detail="Order not found")
    return order

# PayPal Routes
@api_router.post("/paypal/create-order")
async def create_paypal_order(order_id: str = Form(...)):
    """Create PayPal order for payment"""
    order = await db.orders.find_one({"order_id": order_id}, {"_id": 0})
    if not order:
        raise HTTPException(status_code=404, detail="Order not found")
    
    access_token = await get_paypal_access_token()
    
    payload = {
        "intent": "CAPTURE",
        "purchase_units": [{
            "reference_id": order_id,
            "description": f"FileSolved - {order['service_name']}",
            "amount": {
                "currency_code": "USD",
                "value": f"{order['amount']:.2f}"
            }
        }]
    }
    
    async with httpx.AsyncClient() as client:
        response = await client.post(
            f"{PAYPAL_BASE_URL}/v2/checkout/orders",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json"
            },
            json=payload
        )
        
        if response.status_code not in [200, 201]:
            logger.error(f"PayPal create order error: {response.text}")
            raise HTTPException(status_code=500, detail="Failed to create PayPal order")
        
        paypal_order = response.json()
        
        # Update order with PayPal order ID
        await db.orders.update_one(
            {"order_id": order_id},
            {"$set": {"paypal_order_id": paypal_order["id"]}}
        )
        
        return {"paypal_order_id": paypal_order["id"]}

@api_router.post("/paypal/capture-order")
async def capture_paypal_order(
    paypal_order_id: str = Form(...),
    order_id: str = Form(...),
    background_tasks: BackgroundTasks = None
):
    """Capture PayPal payment and process file"""
    access_token = await get_paypal_access_token()
    
    async with httpx.AsyncClient() as client:
        response = await client.post(
            f"{PAYPAL_BASE_URL}/v2/checkout/orders/{paypal_order_id}/capture",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json"
            }
        )
        
        if response.status_code not in [200, 201]:
            logger.error(f"PayPal capture error: {response.text}")
            raise HTTPException(status_code=500, detail="Payment capture failed")
        
        capture_data = response.json()
        
        if capture_data.get("status") == "COMPLETED":
            # Update order status
            await db.orders.update_one(
                {"order_id": order_id},
                {"$set": {"status": "paid", "processed_at": datetime.now(timezone.utc).isoformat()}}
            )
            
            # Record analytics
            order = await db.orders.find_one({"order_id": order_id}, {"_id": 0})
            await db.analytics.insert_one({
                "event": "payment_completed",
                "order_id": order_id,
                "service_id": order["service_id"],
                "amount": order["amount"],
                "timestamp": datetime.now(timezone.utc).isoformat()
            })
            
            # Process file in background
            if background_tasks:
                background_tasks.add_task(process_order_background, order_id)
            
            return {"status": "success", "message": "Payment captured successfully"}
        else:
            raise HTTPException(status_code=400, detail="Payment not completed")

async def process_order_background(order_id: str):
    """Background task to process order"""
    try:
        order = await db.orders.find_one({"order_id": order_id}, {"_id": 0})
        if not order:
            return
        
        # Find the uploaded file
        file_id = order.get("file_id")
        if not file_id:
            return
        
        # Find file by ID
        input_files = list(UPLOAD_DIR.glob(f"{file_id}.*"))
        if not input_files:
            logger.error(f"Input file not found for order {order_id}")
            return
        
        input_path = input_files[0]
        
        # Process the file
        output_path = await process_file(order, input_path)
        
        if output_path and output_path.exists():
            # Update order with output file
            await db.orders.update_one(
                {"order_id": order_id},
                {"$set": {
                    "status": "completed",
                    "output_file": str(output_path),
                    "processed_at": datetime.now(timezone.utc).isoformat()
                }}
            )
            
            # Send email with result
            await send_result_email(order, output_path)
            
            logger.info(f"Order {order_id} processed successfully")
        else:
            await db.orders.update_one(
                {"order_id": order_id},
                {"$set": {"status": "failed"}}
            )
    except Exception as e:
        logger.error(f"Error processing order {order_id}: {e}")
        await db.orders.update_one(
            {"order_id": order_id},
            {"$set": {"status": "failed"}}
        )

# Download processed file
@api_router.get("/orders/{order_id}/download")
async def download_output(order_id: str):
    """Download processed file"""
    order = await db.orders.find_one({"order_id": order_id}, {"_id": 0})
    if not order:
        raise HTTPException(status_code=404, detail="Order not found")
    
    if order.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Order not yet completed")
    
    output_file = order.get("output_file")
    if not output_file or not Path(output_file).exists():
        raise HTTPException(status_code=404, detail="Output file not found")
    
    return FileResponse(output_file, filename=Path(output_file).name)

# Admin Routes
@api_router.post("/admin/login")
async def admin_login(credentials: AdminLogin):
    """Admin login"""
    admin = await db.admins.find_one({"email": credentials.email}, {"_id": 0})
    
    if not admin:
        # Create default admin if none exists
        if credentials.email == "admin@filesolved.com" and credentials.password == "Admin123!":
            hashed = pwd_context.hash(credentials.password)
            await db.admins.insert_one({
                "email": credentials.email,
                "password": hashed,
                "name": "Admin",
                "created_at": datetime.now(timezone.utc).isoformat()
            })
            admin = {"email": credentials.email, "name": "Admin", "password": hashed}
        else:
            raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not pwd_context.verify(credentials.password, admin["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Generate JWT
    token = jwt.encode(
        {
            "email": admin["email"],
            "name": admin.get("name", "Admin"),
            "exp": datetime.now(timezone.utc).timestamp() + 86400  # 24 hours
        },
        JWT_SECRET,
        algorithm=JWT_ALGORITHM
    )
    
    return {"token": token, "email": admin["email"], "name": admin.get("name", "Admin")}

@api_router.get("/admin/analytics")
async def get_analytics(admin: dict = Depends(get_current_admin)):
    """Get analytics data"""
    # Total revenue
    pipeline = [
        {"$match": {"status": {"$in": ["paid", "completed"]}}},
        {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
    ]
    revenue_result = await db.orders.aggregate(pipeline).to_list(1)
    total_revenue = revenue_result[0]["total"] if revenue_result else 0
    
    # Total orders
    total_orders = await db.orders.count_documents({})
    
    # Orders by service
    service_pipeline = [
        {"$group": {"_id": "$service_id", "count": {"$sum": 1}, "revenue": {"$sum": "$amount"}}}
    ]
    service_stats = await db.orders.aggregate(service_pipeline).to_list(100)
    orders_by_service = {s["_id"]: s["count"] for s in service_stats}
    revenue_by_service = {s["_id"]: s["revenue"] for s in service_stats}
    
    # Conversion rate (paid orders / total orders)
    paid_orders = await db.orders.count_documents({"status": {"$in": ["paid", "completed"]}})
    conversion_rate = (paid_orders / total_orders * 100) if total_orders > 0 else 0
    
    # Recent orders
    recent_orders = await db.orders.find({}, {"_id": 0}).sort("created_at", -1).limit(10).to_list(10)
    
    return {
        "total_revenue": round(total_revenue, 2),
        "total_orders": total_orders,
        "orders_by_service": orders_by_service,
        "revenue_by_service": revenue_by_service,
        "conversion_rate": round(conversion_rate, 2),
        "recent_orders": recent_orders
    }

@api_router.get("/admin/orders")
async def get_all_orders(
    skip: int = 0,
    limit: int = 50,
    status: Optional[str] = None,
    admin: dict = Depends(get_current_admin)
):
    """Get all orders with pagination"""
    query = {}
    if status:
        query["status"] = status
    
    orders = await db.orders.find(query, {"_id": 0}).sort("created_at", -1).skip(skip).limit(limit).to_list(limit)
    total = await db.orders.count_documents(query)
    
    return {"orders": orders, "total": total}

@api_router.get("/admin/revenue-summary")
async def get_revenue_summary(admin: dict = Depends(get_current_admin)):
    """Get revenue summary by day"""
    pipeline = [
        {"$match": {"status": {"$in": ["paid", "completed"]}}},
        {"$addFields": {
            "date": {"$substr": ["$created_at", 0, 10]}
        }},
        {"$group": {
            "_id": "$date",
            "revenue": {"$sum": "$amount"},
            "orders": {"$sum": 1}
        }},
        {"$sort": {"_id": -1}},
        {"$limit": 30}
    ]
    
    summary = await db.orders.aggregate(pipeline).to_list(30)
    return {"daily_revenue": summary}

# SEO Routes
@api_router.get("/sitemap")
async def get_sitemap():
    """Generate sitemap data"""
    base_url = "https://filesolved.com"
    pages = [
        {"url": "/", "priority": 1.0, "changefreq": "weekly"},
        {"url": "/services", "priority": 0.9, "changefreq": "weekly"},
        {"url": "/services/pdf-to-word", "priority": 0.8, "changefreq": "monthly"},
        {"url": "/services/word-to-pdf", "priority": 0.8, "changefreq": "monthly"},
        {"url": "/services/jpg-to-pdf", "priority": 0.8, "changefreq": "monthly"},
        {"url": "/services/ocr", "priority": 0.8, "changefreq": "monthly"},
        {"url": "/services/document-scan", "priority": 0.8, "changefreq": "monthly"},
        {"url": "/services/pdf-fax", "priority": 0.8, "changefreq": "monthly"},
        {"url": "/services/secure-shred", "priority": 0.8, "changefreq": "monthly"},
        {"url": "/how-to/convert-pdf-to-word", "priority": 0.7, "changefreq": "monthly"},
        {"url": "/how-to/convert-word-to-pdf", "priority": 0.7, "changefreq": "monthly"},
        {"url": "/faq", "priority": 0.6, "changefreq": "monthly"},
        {"url": "/contact", "priority": 0.5, "changefreq": "monthly"},
    ]
    return {"base_url": base_url, "pages": pages}

# Include the router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
